import os
import logging
from typing import List, Dict, Any, Optional
from pathlib import Path
import asyncio
from datetime import datetime

import openai
from sqlalchemy.orm import Session
from sqlalchemy import text
import numpy as np
from sentence_transformers import SentenceTransformer
import PyPDF2
import docx
from io import BytesIO

from ..database import get_db
from ..models import ChatLog, UsageStats
from ..models import Document, DocumentChunk
from dotenv import load_dotenv 

load_dotenv() 
logger = logging.getLogger(__name__)

class RAGService:
    def __init__(self):
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise ValueError("OPENAI_API_KEY not found in environment variables")
        self.openai_client = openai.OpenAI(api_key=api_key)
        self.embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        self.chunk_size = 1000
        self.chunk_overlap = 200
        
    async def ingest_document(self, file_path: str, file_name: str, file_type: str, db: Session) -> Document:
        """Ingest a document, extract text, create embeddings, and store in database."""
        try:
            # Extract text from document
            text_content = await self._extract_text(file_path, file_type)
            
            # Create document record
            document = Document(
                filename=file_name,
                file_type=file_type,
                file_path=file_path,
                content=text_content,
                status="processing",
                created_at=datetime.utcnow()
            )
            db.add(document)
            db.commit()
            db.refresh(document)
            
            # Create chunks
            chunks = self._create_chunks(text_content)
            
            # Generate embeddings and store chunks
            for i, chunk_text in enumerate(chunks):
                embedding = await self._generate_embedding(chunk_text)
                
                chunk = DocumentChunk(
                    document_id=document.id,
                    chunk_index=i,
                    content=chunk_text,
                    embedding=embedding.tolist(),
                    created_at=datetime.utcnow()
                )
                db.add(chunk)
            
            # Update document status
            document.status = "completed"
            document.processed_at = datetime.utcnow()
            db.commit()
            
            logger.info(f"Successfully ingested document: {file_name}")
            return document
            
        except Exception as e:
            logger.error(f"Error ingesting document {file_name}: {str(e)}")
            if 'document' in locals():
                document.status = "failed"
                document.error_message = str(e)
                db.commit()
            raise
    
    async def _extract_text(self, file_path: str, file_type: str) -> str:
        """Extract text content from various file types."""
        try:
            if file_type.lower() == 'pdf':
                return await self._extract_pdf_text(file_path)
            elif file_type.lower() in ['doc', 'docx']:
                return await self._extract_docx_text(file_path)
            elif file_type.lower() == 'txt':
                with open(file_path, 'r', encoding='utf-8') as file:
                    return file.read()
            else:
                raise ValueError(f"Unsupported file type: {file_type}")
        except Exception as e:
            logger.error(f"Error extracting text from {file_path}: {str(e)}")
            raise
    
    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file."""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    
    async def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        doc = docx.Document(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    def _create_chunks(self, text: str) -> List[str]:
        """Split text into overlapping chunks."""
        chunks = []
        start = 0
        
        while start < len(text):
            end = start + self.chunk_size
            chunk = text[start:end]
            
            # Try to break at sentence boundary
            if end < len(text):
                last_period = chunk.rfind('.')
                last_newline = chunk.rfind('\n')
                break_point = max(last_period, last_newline)
                
                if break_point > start + self.chunk_size // 2:
                    chunk = text[start:start + break_point + 1]
                    end = start + break_point + 1
            
            chunks.append(chunk.strip())
            start = end - self.chunk_overlap
            
        return [chunk for chunk in chunks if chunk.strip()]
    
    async def _generate_embedding(self, text: str) -> np.ndarray:
        """Generate embedding for text using sentence transformer."""
        try:
            # Use local embedding model for faster processing
            embedding = self.embedding_model.encode(text)
            return embedding
        except Exception as e:
            logger.error(f"Error generating embedding: {str(e)}")
            # Fallback to OpenAI embeddings
            response = await self.openai_client.embeddings.create(
                model="text-embedding-ada-002",
                input=text
            )
            return np.array(response.data[0].embedding)
    
    async def retrieve_relevant_chunks(self, query: str, db: Session, top_k: int = 5) -> List[Dict[str, Any]]:
        """Retrieve most relevant document chunks for a query."""
        try:
            # Generate query embedding
            query_embedding = await self._generate_embedding(query)
            
            # Use vector similarity search (assuming PostgreSQL with pgvector)
            # This is a simplified version - in production, use proper vector search
            chunks = db.query(DocumentChunk).all()
            
            similarities = []
            for chunk in chunks:
                chunk_embedding = np.array(chunk.embedding)
                similarity = np.dot(query_embedding, chunk_embedding) / (
                    np.linalg.norm(query_embedding) * np.linalg.norm(chunk_embedding)
                )
                similarities.append({
                    'chunk': chunk,
                    'similarity': similarity,
                    'document': chunk.document
                })
            
            # Sort by similarity and return top_k
            similarities.sort(key=lambda x: x['similarity'], reverse=True)
            
            return similarities[:top_k]
            
        except Exception as e:
            logger.error(f"Error retrieving relevant chunks: {str(e)}")
            return []
    
    async def generate_rag_response(self, query: str, context_chunks: List[Dict[str, Any]]) -> str:
        """Generate response using RAG with retrieved context."""
        try:
            # Prepare context from retrieved chunks
            context = "\n\n".join([
                f"Document: {chunk['document'].filename}\n{chunk['chunk'].content}"
                for chunk in context_chunks
            ])
            
            # Create prompt with context
            prompt = f"""You are a helpful healthcare support assistant for a medical equipment company. 
Use the following context to answer the user's question. If the context doesn't contain relevant information, 
say so and provide general guidance.

Context:
{context}

Question: {query}

Answer:"""
            
            response = await self.openai_client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": "You are a helpful healthcare support assistant."},
                    {"role": "user", "content": prompt}
                ],
                max_tokens=500,
                temperature=0.7
            )
            
            return response.choices[0].message.content
            
        except Exception as e:
            logger.error(f"Error generating RAG response: {str(e)}")
            return "I apologize, but I'm having trouble accessing the relevant information right now. Please try again or contact support."